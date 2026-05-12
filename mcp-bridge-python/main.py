import pandas as pd
import fitz  # PyMuPDF for PDFs
from docx import Document # For Word docs
import subprocess
import os
import json
from mcp.server.fastmcp import FastMCP

# Initialize the server with a recruiter-friendly name
mcp = FastMCP("Data-Intelligence-Bridge")

@mcp.tool()
def analyze_csv(file_path: str) -> str:
    """Reads a CSV and returns a statistical summary and missing value audit."""
    try:
        path = os.path.expanduser(file_path)
        df = pd.read_csv(path)
        summary = {
            "rows": len(df),
            "columns": list(df.columns),
            "missing_data": df.isnull().sum().to_dict(),
            "stats": df.describe().to_dict() if not df.select_dtypes(include='number').empty else "No numeric columns"
        }
        return f"### CSV Audit Report for {os.path.basename(path)}\n{json.dumps(summary, indent=2)}"
    except Exception as e:
        return f"CSV Error: {str(e)}"

@mcp.tool()
def analyze_pdf(file_path: str) -> str:
    """Extracts text and metadata from a PDF for content auditing."""
    try:
        path = os.path.expanduser(file_path)
        doc = fitz.open(path)
        text = "".join([page.get_text() for page in doc.pages(0, min(3, doc.page_count))])
        return f"### PDF Analysis: {os.path.basename(path)}\n**Metadata:** {doc.metadata}\n**Preview:** {text[:800]}..."
    except Exception as e:
        return f"PDF Error: {str(e)}"

@mcp.tool()
def analyze_docx(file_path: str) -> str:
    """Parses a Word document and returns structural insights."""
    try:
        path = os.path.expanduser(file_path)
        doc = Document(path)
        preview = "\n".join([p.text for p in doc.paragraphs[:10]])
        return f"### Word Doc Analysis: {os.path.basename(path)}\n**Paragraphs:** {len(doc.paragraphs)}\n**Preview:** {preview[:800]}..."
    except Exception as e:
        return f"Word Error: {str(e)}"

@mcp.tool()
def open_system_app(app_name: str) -> str:
    """Launches any Mac application (e.g., 'Numbers', 'Preview', 'Safari')."""
    try:
        subprocess.run(["open", "-a", app_name], check=True)
        return f"Success: Launched {app_name}."
    except Exception:
        return f"Error: Could not find application '{app_name}'."

if __name__ == "__main__":
    mcp.run()